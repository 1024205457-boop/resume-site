#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = '楷体'
FONT_EN = 'Times New Roman'

def set_font(run, size=10, bold=False):
    run.font.size = Pt(size)
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    run.bold = bold

def remove_cell_borders(cell):
    """去掉单元格所有边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def remove_table_borders(table):
    """去掉表格自身边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    tblPr.append(borders)

def set_cell_margin(cell, top=0, bottom=0, left=0, right=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def _make_three_col_table(doc, left, center, right, col_widths, size=10, left_bold=True):
    """通用无边框三列表格"""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    for i, w in enumerate(col_widths):
        table.columns[i].width = Cm(w)

    cells = table.rows[0].cells
    for c in cells:
        remove_cell_borders(c)
        set_cell_margin(c, top=0, bottom=0, left=0, right=0)
        c.paragraphs[0].paragraph_format.space_before = Pt(1)
        c.paragraphs[0].paragraph_format.space_after = Pt(1)

    r = cells[0].paragraphs[0].add_run(left)
    set_font(r, size=size, bold=left_bold)
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    r = cells[1].paragraphs[0].add_run(center)
    set_font(r, size=size, bold=False)
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = cells[2].paragraphs[0].add_run(right)
    set_font(r, size=size, bold=False)
    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return table

def add_three_col_table(doc, left, center, right, size=10, left_bold=True):
    """实习/项目经历用：左列宽一些放公司/项目名"""
    return _make_three_col_table(doc, left, center, right, [7, 6, 4.4], size, left_bold)

def add_edu_table(doc, left, center, right, size=10):
    """教育经历用：中间列更宽放学院+专业"""
    return _make_three_col_table(doc, left, center, right, [3.2, 11, 3.2], size, left_bold=True)

def add_section_title(doc, text):
    """段落标题 + 底部横线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, bold=True)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_text(doc, text, size=9.5, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def add_bullet(doc, label, content, size=9.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(label + '：')
    set_font(r1, size=size, bold=True)
    r2 = p.add_run(content)
    set_font(r2, size=size, bold=False)
    return p

# ========== 生成简历 ==========
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# ===== 姓名 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('姜慧男')
set_font(run, size=16, bold=True)

# ===== 联系方式 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run('时间：ASAP, 6个月及以上 | 电话&微信：15241569480 | 邮箱：jianghuinan@connect.hku.hk')
set_font(run, size=9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run('领英：https://www.linkedin.com/in/lily-jiang-325913325')
set_font(run, size=9)

# ===== 教育经历 =====
add_section_title(doc, '教育经历')

add_edu_table(doc, '香港大学', '商学院 市场学硕士 量化营销方向', '2024.9 - 2026.12')
add_text(doc, 'GPA：3.8/4.3   荣誉：学业优秀奖学金（36,000港币）')

add_edu_table(doc, '厦门大学', '管理学院 市场营销&人类学（双学位）本科 社会心理学方向', '2020.9 - 2024.6')
add_text(doc, 'GPA：3.6/4.0   荣誉：优秀学生干部、"挑战杯"省银奖、"互联网+"校金奖、4项校级奖学金')

# ===== 个人技能 =====
add_section_title(doc, '个人技能')

add_bullet(doc, 'AI工具与开发', 'Claude Code, Cursor, LangChain, RAG, Prompt Engineering, 智谱GLM API, Coze, 腾讯元器')
add_bullet(doc, '产品能力', '用户调研, 需求分析, PRD撰写, Figma原型设计, X-mind流程梳理, A/B Test')
add_bullet(doc, '数据与技术', 'Python, SQL, SPSS, Flask, Next.js, D3.js, 数据看板搭建')
add_bullet(doc, '运营与传播', '小红书（单篇56w+浏览）, 抖音（单条1w+浏览）, 腾讯产品经理训练营')
add_bullet(doc, '语言能力', '英语（雅思7）')

# ===== 实习经历 =====
add_section_title(doc, '实习经历')

add_three_col_table(doc, '猿辅导（斑马百科）', '用户增长AI运营实习生', '2026.2 - 2026.3')

add_bullet(doc, '产品设计',
    '针对团队周报数据处理耗时2-3h的问题，通过用户调研梳理4大业务模块15个功能的操作流程与数据流向，完成需求拆解、信息架构设计与交互原型输出，落地数据处理系统将流程缩短至分钟级，推动全组采纳使用。')

add_bullet(doc, 'AI应用落地',
    '为替代每周3人的私域文案与素材人工产出，设计Prompt策略（主题循环、时令感知、历史去重），接入智谱大模型API实现7天×3渠道共63条文案自动生产；对接魔袋AI图生图批量生成导购卡与配图，打通"生成→审核→推送"全链路，人力从3人降至0人。')

add_bullet(doc, '数据分析',
    '搭建多数据源（Pipe、Mario、Dora、BI）自动化采集与清洗流程，设计转化漏斗（加微UV→领取UV→激活UV→购买）与流失分析模型，输出标准化周报表格支撑运营团队决策。')

# ===== 项目经历 =====
add_section_title(doc, '项目经历')

# --- Venn AI ---
add_three_col_table(doc, 'Venn AI 维恩图智能笔记', '产品经理 & 全栈开发', '2026.3')

add_bullet(doc, '产品设计',
    '针对思维导图只能表达树状层级、无法呈现概念交叉关系的痛点，定义"收集+整理"双模式产品架构（关键词由AI展开子概念与交叉关系/上传笔记忠于原文提取结构），设计2/3层嵌套生成策略与手动补充机制，已上线覆盖学科对比、考试复习、竞品分析等7大场景。')

add_bullet(doc, 'AI能力设计',
    '接入智谱GLM大模型API，针对双模式分别设计Prompt策略（收集模式允许AI知识扩展，整理模式约束严格忠于原文），实现自然语言到结构化维恩图JSON的端到端生成，支持.txt/.md/.docx多格式文件上传（最大5万字）。')

add_bullet(doc, '全栈开发',
    '采用Next.js + D3.js + SQLite技术栈，独立开发可缩放、可编辑的交互式嵌套维恩图组件（简洁/透视双视图、节点增删改、深入分析二次展开），完成部署上线并开源至GitHub。')

# --- 抖音黑客松 ---
add_three_col_table(doc, '抖音AI创变者跨年黑客松"优秀产品"', '产品经理 & 全栈开发', '2025.12 - 2026.1')

add_bullet(doc, '产品定义',
    '通过竞品分析和行业调研找准"情绪卸载"切入点，定义"多模态AI情绪识别与发泄APP"产品定位；用Figma设计原型，需求变更率控制在10%以内，路演当日获2家孵化器支持，200+观众中69人愿意试用。')

add_bullet(doc, 'AI应用开发',
    '利用LangChain框架重写后端，优化memory存储调用逻辑并调试引导问题生成Prompt；前端通过Kimi设计交互动画，完成APP核心功能开发与部署。')

# --- 腾讯云 ---
add_three_col_table(doc, '腾讯云Agent Mini-hackathon"三等奖"', '产品经理 & 全栈开发', '2025.12')

add_bullet(doc, '产品定义',
    '通过行业报告归纳"海外护肤品成分核验难"痛点，洞察71%成分党的消费潜力，定义"敏敏肌选品"导购Agent产品定位，梳理了解-搜索-电商操作闭环并编写PRD。')

add_bullet(doc, 'AI Agent开发',
    '在腾讯元器平台基于RAG机制设计Prompt与知识库，调整API对接策略，7小时内完成从产品设计到微信小程序上线的全流程交付。')

# --- 清华 ---
add_three_col_table(doc, '清华vibe coding黑客松"最佳创意"', '产品经理', '2025.12')

add_bullet(doc, '产品设计与落地',
    '通过4场用户访谈精准定义"线上会议难以优雅离场"痛点，参考微信弹窗设计高仿真来电模拟软件"真的很忙"，3小时内用Cursor完成开发上线，目前拥有300+忠实用户，小红书获6w+用户关注。')

# ===== 保存 =====
output = '/Users/kanyun/Documents/简历-姜慧男-AI产品经理.docx'
doc.save(output)
print(f'已生成: {output}')
