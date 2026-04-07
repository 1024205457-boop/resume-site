#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 页面边距
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(0.3)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.0

def add_name(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)

def add_contact(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(80, 80, 80)

def add_section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    # 下划线效果
    from docx.oxml import OxmlElement
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_exp_header(company, title, time, col_widths=None):
    from docx.oxml import OxmlElement
    # 默认列宽：中间3000居中，左右各3603
    if col_widths is None:
        col_widths = [3603, 3000, 3603]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    # 隐藏边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'none')
        e.set(qn('w:sz'), '0')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), 'auto')
        borders.append(e)
    tblPr.append(borders)
    # 去掉表格前后间距
    spacing = OxmlElement('w:tblCellSpacing')
    spacing.set(qn('w:w'), '0')
    spacing.set(qn('w:type'), 'dxa')
    tblPr.append(spacing)
    tblGrid = tbl.tblGrid
    for i, w in enumerate(col_widths):
        tc = table.rows[0].cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(w))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
    # 去掉单元格内边距
    for i in range(3):
        tc = table.rows[0].cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side in ('top', 'bottom'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:w'), '0')
            el.set(qn('w:type'), 'dxa')
            tcMar.append(el)
        tcPr.append(tcMar)
    # 左：公司
    c0 = table.rows[0].cells[0].paragraphs[0]
    c0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    c0.paragraph_format.space_before = Pt(2)
    c0.paragraph_format.space_after = Pt(0)
    r1 = c0.add_run(company)
    r1.bold = True
    r1.font.size = Pt(10)
    # 中：岗位
    c1 = table.rows[0].cells[1].paragraphs[0]
    c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c1.paragraph_format.space_before = Pt(2)
    c1.paragraph_format.space_after = Pt(0)
    r2 = c1.add_run(title)
    r2.bold = True
    r2.font.size = Pt(10)
    # 右：时间
    c2 = table.rows[0].cells[2].paragraphs[0]
    c2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    c2.paragraph_format.space_before = Pt(2)
    c2.paragraph_format.space_after = Pt(0)
    r3 = c2.add_run(time)
    r3.font.size = Pt(10)

def add_bg(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(80, 80, 80)

def add_bullet(label, content):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.4)
    r0 = p.add_run('· ')
    r0.font.size = Pt(10)
    r1 = p.add_run(label + '：')
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(content)
    r2.font.size = Pt(10)

def add_line(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(10)

# ========== 正文 ==========

add_name('姜慧男')
add_contact('时间：ASAP, 6个月及以上 | 电话&微信：15241569480 | 邮箱：jianghuinan@connect.hku.hk')
add_contact('领英：linkedin.com/in/lily-jiang-325913325 | 岗位：产品实习生')

# 教育经历
add_section_title('教育经历')

add_exp_header('香港大学  商学院', '市场学硕士（量化营销方向）', '2024.9 - 2026.12', [2800, 4606, 2800])
add_line('GPA：3.8/4.3　　学业优秀奖学金（36,000港币）')

add_exp_header('厦门大学  管理学院', '市场营销&人类学双学位（社会心理学方向）', '2020.9 - 2024.6', [2800, 4606, 2800])
add_line('GPA：3.6/4.0　　优秀学生干部、"挑战杯"省银奖、"互联网+"校金奖、4项校级奖学金')

# 个人技能
add_section_title('个人技能')

add_bullet('AI工具与开发', 'Claude Code, Cursor, LangChain, RAG, Prompt Engineering, 智谱GLM API, Coze, 腾讯元器')
add_bullet('产品能力', '用户调研, 需求分析, PRD撰写, Figma原型设计, X-mind流程梳理, A/B Test')
add_bullet('数据与技术', 'Python, SQL, SPSS, Flask, Next.js, D3.js, 数据看板搭建')
add_bullet('运营与传播', '小红书（单篇56w+浏览）, 抖音（单条1w+浏览）, 腾讯产品经理训练营')
add_bullet('语言能力', '英语（雅思7）')

# 实习经历
add_section_title('实习经历')

add_exp_header('猿辅导（斑马百科）', '用户增长AI运营实习生', '2026.2 - 至今')
add_bg('实习背景：利用Claude Code和Cursor独立编写并部署AI自动化工作流，服务于用户增长团队日常运营。')
add_bullet('需求规划', '设计AI基建需求平台，定义6大能力域框架（数据读写、平台操作、消息交互、内容生成、流程串联），收集4条业务线7个自动化需求并完成需求翻译，独立推进从需求到上线的全流程。')
add_bullet('工具开发', '独立开发5套AI自动化工具并部署至服务器配置进程守护，覆盖会员触达、数据处理、推送全链路，团队日均重复操作耗时从2h降至10min，人效提升12倍。')
add_bullet('数据搭建', '搭建获客概览数据看板（Flask+pandas），覆盖4大业务模块、5大渠道，支持Excel自动生成透视表与数据下载，追踪21日转化率、eLTV与CAC。')

add_exp_header('意略明咨询（艾昆玮全资子公司）', '电商策略分析实习生', '2025.4 - 2025.8')
add_bg('实习背景：负责惠氏、伊利、强生安视优的社媒触点到电商的转化效率与策略分析。')
add_bullet('用户洞察', '通过800份问卷、12段焦点小组访谈、67,000条行为数据，提炼痛点与改进建议。')
add_bullet('漏斗分析', '搭建品牌流转漏斗模型，追踪曝光→浏览→购买转化率，还原社媒到电商决策路径，节省15%投放。')

# 项目经历
add_section_title('项目经历')

add_exp_header('抖音AI创变者计划跨年黑客松"优秀产品"', '产品经理&全栈开发', '2025.12 - 2026.1')
add_bullet('竞品定位', '深入研究竞品功能找准"情绪卸载"切入点，提出多模态AI情绪识别APP定位；路演当日获2家孵化器支持，200余名观众中69人愿意试用。')
add_bullet('产品落地', '通过Figma设计产品原型、X-mind梳理业务流程，需求变更率控制在10%以内；利用langchain重写后端逻辑并完成部署上线。')

add_exp_header('腾讯云Agent Mini-hackathon"三等奖"', '产品经理&全栈开发', '2025.12')
add_bullet('需求洞察', '归纳"海外护肤品成分核验难"用户痛点，通过数据分析洞察71%成分党消费潜力，明确客服导购Agent产品方向与核心功能。')
add_bullet('快速交付', '基于RAG检索增强机制优化prompt设计与API调用链路，7小时内独立完成微信小程序从产品设计、开发到上线的全流程交付。')

add_exp_header('清华大学vibe coding黑客松"最佳创意"', '产品经理', '2025.12')
add_bullet('产品设计', '通过4场用户访谈定义"线上会议难以优雅离场"痛点，设计来电模拟软件"真的很忙"，3小时内用Cursor开发上线，累计300+忠实用户，小红书6w+关注。')

add_exp_header('"腾讯公益"标语A/B Test', '运营分析', '2023.12 - 2024.4')
add_bullet('实验优化', '基于1645篇文献设计A/B Test优化募捐文案，捐款金额+33%。')

output = '/Users/kanyun/Documents/简历-姜慧男-AI产品经理.docx'
doc.save(output)
print(f'已生成: {output}')
